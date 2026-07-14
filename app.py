import streamlit as st
import pandas as pd
import json
import os
import requests
import re
from docx import Document
import io
import html
import difflib

# 1. Page Config & Setup
st.set_page_config(
    page_title="博幼會議專案追蹤系統",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Load theme state - Default is now Light mode (白天模式)
if "theme" not in st.session_state:
    st.session_state.theme = "light"

IS_DARK = st.session_state.theme == "dark"

# Theme Variables (Optimized for premium Light Mode with high-contrast text and labels)
BG_COLOR = "#0B0F19" if IS_DARK else "#F1F5F9"
TEXT_COLOR = "#E2E8F0" if IS_DARK else "#1E293B"
TITLE_COLOR = "#FFFFFF" if IS_DARK else "#0F172A"
CARD_BG = "rgba(17, 24, 39, 0.75)" if IS_DARK else "#FFFFFF"
BORDER_COLOR = "rgba(255, 255, 255, 0.08)" if IS_DARK else "#CBD5E1"
TEXT_MUTED = "#94A3B8" if IS_DARK else "#475569"
CARD_PROGRESS_BG = "rgba(16, 185, 129, 0.03)" if IS_DARK else "rgba(16, 185, 129, 0.04)"
CARD_PROGRESS_BORDER = "rgba(16, 185, 129, 0.15)" if IS_DARK else "rgba(16, 185, 129, 0.25)"

# --- Parser & Router Helpers ---
DEPT_KEYWORDS = {
    "教材研發組": ["教材", "教材組", "研發編輯", "英文編輯", "數學編輯", "閱讀編輯", "美編", "課輔教材", "檢測卷", "能力檢定", "科普", "程式", "心算"],
    "社工特教組": ["社工", "特教", "特殊教育", "個督", "急難救助", "安心基金", "特教生", "早療", "小手計畫", "輔導知能", "輔導工作", "韋志", "己智", "純菁", "谷燕婕"],
    "學區營運組": ["學區", "營運", "營運組", "學區營運", "課師", "鐘點", "師培", "週誌", "週誌優化", "成績系統", "開課系統", "培訓課程", "秉謙", "賢明", "睿豪"],
    "畢業生組": ["畢業生", "畢業生組", "學生會", "獎學金", "獎助學金", "就學就業", "補習", "大立光", "建銘", "綉琴", "聖仁", "會內獎學金", "校訪", "中離", "書婷", "昆良"],
    "處長室/行政管理": ["處長", "行政管理", "整併", "人服", "出勤", "加班", "外訓", "認列", "品管", "品圈", "工作規則", "工作細則", "合約", "顧問", "杜瀛", "桓根", "執行長"]
}

def detect_department(text):
    matched_dept = "處長室/行政管理"
    max_score = 0
    for dept, kw_list in DEPT_KEYWORDS.items():
        score = 0
        for kw in kw_list:
            if kw.lower() in text.lower():
                score += 1
        if score > max_score:
            max_score = score
            matched_dept = dept
    return matched_dept

def extract_date_from_filename(filename):
    """Automatically extract Minguo date from filename (e.g. 1150616 -> 2026.06.16)."""
    match = re.search(r"(\d{6,7})", filename)
    if match:
        num_str = match.group(1)
        try:
            if len(num_str) == 7: # e.g. 1150616 -> 115 + 1911 = 2026
                year = int(num_str[:3]) + 1911
                month = num_str[3:5]
                day = num_str[5:]
                return f"{year}.{month}.{day}"
            elif len(num_str) == 6: # e.g. 990616 -> 99 + 1911 = 2010
                year = int(num_str[:2]) + 1911
                month = num_str[2:4]
                day = num_str[4:]
                return f"{year}.{month}.{day}"
        except:
            pass
    return pd.Timestamp.now().strftime("%Y.%m.%d")

def find_fuzzy_match(parsed_task, existing_tasks):
    """Find the best fuzzy matching existing task based on both title and content similarity."""
    best_match = None
    best_score = 0.0
    best_reason = ""
    
    parsed_title = parsed_task.get("title", "").strip().lower()
    parsed_content = parsed_task.get("content", "").strip().lower()
    
    for et in existing_tasks:
        et_title = et.get("title", "").strip().lower()
        et_content = et.get("content", "").strip().lower()
        
        # Exact title check (100% match)
        if parsed_title == et_title:
            return et, 1.0, "任務名稱完全一致"
            
        t_score = difflib.SequenceMatcher(None, parsed_title, et_title).ratio()
        c_score = difflib.SequenceMatcher(None, parsed_content, et_content).ratio()
        
        # Take the maximum similarity score between title and content
        max_score = max(t_score, c_score)
        if max_score > best_score:
            best_score = max_score
            best_match = et
            if t_score >= c_score:
                best_reason = f"任務名稱相似度 {round(t_score * 100)}%"
            else:
                best_reason = f"決議說明相似度 {round(c_score * 100)}%"
                
    if best_score >= 0.70: # Threshold of 70%
        return best_match, best_score, best_reason
    return None, 0.0, ""

def parse_docx_to_tasks(file_bytes, filename):
    doc = Document(io.BytesIO(file_bytes))
    meeting_date = extract_date_from_filename(filename)
    
    # Check if Gemini API Key is available
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key and hasattr(st, "secrets") and "GEMINI_API_KEY" in st.secrets:
        api_key = st.secrets["GEMINI_API_KEY"]
        
    if api_key:
        try:
            # Extract all text and table contents from docx
            full_text = []
            for p in doc.paragraphs:
                t = p.text.strip()
                if t:
                    full_text.append(t)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        full_text.append(row_text)
            doc_content = "\n".join(full_text)
            
            # Call Gemini API
            url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={api_key}"
            headers = {"Content-Type": "application/json"}
            prompt = (
                "你是一個博幼社會福利基金會的專業行政幕僚助理。請從以下會議記錄文本中，提取出所有需要追蹤的「任務/代辦事項」。\n"
                "請仔細閱讀所有段落與表格。\n\n"
                "【提取與分類規則】\n"
                "1. 任務名稱（title）：簡潔有力，代表任務的主題（不超過80個字）。\n"
                "2. 主責部門（dept）：必須從以下組別中選擇一個最合適的主責組別：\n"
                "   - \"教材研發組\"\n"
                "   - \"社工特教組\"\n"
                "   - \"學區運營組\"\n"
                "   - \"畢業生組\"\n"
                "   - \"處長室/行政管理\"\n"
                "3. 主責對口（owner）：負責該任務的人員名稱，若在文本中沒有提到明確姓名，則填寫「待指派」。\n"
                "4. 任務狀態（status）：若文本中提到「已完成」、「100%」等，狀態設為 \"completed\"，否則設為 \"in_progress\" 或 \"pending\"。\n"
                "5. 優先權（priority）：根據文字描述的緊急程度，歸類為 \"high\"（高）、\"medium\"（中）或 \"low\"（低）。若無特別提及，預設為 \"medium\"。\n"
                "6. 決議細節（content）：該任務的詳細說明與會議決議，請保持完整與準確。\n"
                "7. 目前進度（progress）：任務的最新進度說明。若文本中無特別說明，預設為「自會議記錄匯入」。\n"
                "8. 跨部門協作（is_cross_dept）：這是一個布林值（true/false）。如果任務內容包含跨部門合作、協辦、跨組別、多個部門共同負責等特徵，請設為 true，否則設為 false。\n\n"
                "會議記錄文本如下：\n"
                "\"\"\"\n"
                f"{doc_content}\n"
                "\"\"\""
            )
            
            payload = {
                "contents": [{
                    "parts": [{
                        "text": prompt
                    }]
                }],
                "generationConfig": {
                    "responseMimeType": "application/json",
                    "responseSchema": {
                        "type": "OBJECT",
                        "properties": {
                            "tasks": {
                                "type": "ARRAY",
                                "items": {
                                    "type": "OBJECT",
                                    "properties": {
                                        "title": { "type": "STRING" },
                                        "dept": { 
                                            "type": "STRING", 
                                            "enum": ["教材研發組", "社工特教組", "學區運營組", "畢業生組", "處長室/行政管理"]
                                        },
                                        "owner": { "type": "STRING" },
                                        "status": { 
                                            "type": "STRING", 
                                            "enum": ["pending", "in_progress", "completed"] 
                                        },
                                        "priority": { 
                                            "type": "STRING", 
                                            "enum": ["high", "medium", "low"] 
                                        },
                                        "content": { "type": "STRING" },
                                        "progress": { "type": "STRING" },
                                        "is_cross_dept": { "type": "BOOLEAN" }
                                    },
                                    "required": ["title", "dept", "owner", "status", "priority", "content", "progress", "is_cross_dept"]
                                }
                            }
                        },
                        "required": ["tasks"]
                    }
                }
            }
            
            response = requests.post(url, json=payload, headers=headers, timeout=25)
            if response.status_code == 200:
                data = response.json()
                text_response = data["candidates"][0]["content"]["parts"][0]["text"]
                raw_tasks = json.loads(text_response).get("tasks", [])
                
                # Fill in Python-specific fields
                parsed = []
                for i, t in enumerate(raw_tasks):
                    parsed.append({
                        "id": f"parsed-{int(pd.Timestamp.now().timestamp())}-ai-{i}",
                        "title": t["title"],
                        "dept": t["dept"],
                        "owner": t["owner"],
                        "meeting": filename.replace(".docx", ""),
                        "date": meeting_date,
                        "status": t["status"],
                        "priority": t["priority"],
                        "content": t["content"],
                        "progress": t["progress"],
                        "is_cross_dept": t["is_cross_dept"]
                    })
                st.toast("🔮 已使用 Gemini API 進行 AI 智慧語意解析！", icon="🔮")
                return parsed
        except Exception as e:
            pass
            
    st.toast("❌ 無法使用 Gemini 進行智慧解析，請手動新增任務。", icon="❌")
    return []

def html_escape(text):
    """Prevent XSS and HTML layout breakage in markdown cards."""
    if not isinstance(text, str):
        return str(text) if text is not None else ""
    return html.escape(text)

def get_task_start_date(t):
    """Calculate the project start date based on the earliest date in meeting or progress history."""
    dates = []
    if t.get("created_date"):
        dates.append(t["created_date"])
    for mh in t.get("meeting_history", []):
        if mh.get("date"):
            dates.append(mh["date"])
    for ph in t.get("progress_history", []):
        if ph.get("date"):
            dates.append(ph["date"])
    if t.get("date"):
        dates.append(t["date"])
    if dates:
        try:
            # Sort chronologically YYYY.MM.DD
            return min(dates)
        except:
            pass
    return pd.Timestamp.now().strftime("%Y.%m.%d")

def extract_deadline_flag(text):
    """Scan text for date deadlines (e.g. '6/30前', '7月15日前完成') and return a formatted checkpoint flag."""
    if not text:
        return ""
    # Chinese deadline keywords & date patterns
    patterns = [
        r"(\d{1,2}\s*月\s*\d{1,2}\s*日\s*[前內完截止最晚]*)",
        r"(\d{1,2}\s*/\s*\d{1,2}\s*[前內完截止最晚]*)",
        r"([截止期限預計]*[：:]?\s*\d{1,2}\s*[月/]\s*\d{1,2}\s*日?)",
        r"([明後]天前完成|下週[一二三四五六]前完成|本月[底中]前完成)"
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            val = match.group(1).strip()
            # Verify context to avoid false positives (e.g. simple date references)
            context_match = False
            for kw in ["前", "完", "截止", "期限", "底", "天", "週", "預計", "最晚"]:
                if kw in val:
                    context_match = True
                    break
            if not context_match:
                start = max(0, match.start() - 12)
                end = min(len(text), match.end() + 12)
                surrounding = text[start:end]
                for kw in ["前", "完", "截止", "期限", "底", "天", "週", "預計", "最晚"]:
                    if kw in surrounding:
                        context_match = True
                        break
            if context_match:
                return val
    return ""

def clean_html(html_str):
    """Remove leading whitespace/indentation from multiline HTML blocks to prevent Streamlit rendering them as Markdown code blocks."""
    if not html_str:
        return ""
    return "\n".join([line.strip() for line in html_str.split("\n")])

# Inject Custom CSS for premium glassmorphism layout
st.markdown(clean_html(f"""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');
    
    :root {{
        --bg: {BG_COLOR};
        --text: {TEXT_COLOR};
        --title: {TITLE_COLOR};
        --text-muted: {TEXT_MUTED};
        --card-bg: {CARD_BG};
        --border-color: {BORDER_COLOR};
        --progress-bg: {CARD_PROGRESS_BG};
        --progress-border: {CARD_PROGRESS_BORDER};
    }}

    html, body, [data-testid="stAppViewContainer"], [data-testid="stApp"], .main, .block-container, section[data-testid="stMain"] {{
        background-color: var(--bg) !important;
        color: var(--text) !important;
        font-family: 'Inter', -apple-system, sans-serif !important;
    }}
    
    /* Hide default streamlit headers & buttons */
    header[data-testid="stHeader"], .stDeployButton, footer, [data-testid="stToolbar"] {{
        display: none !important;
    }}
    
    .block-container {{
        padding: 2rem 2.5rem 3rem !important;
        max-width: 1400px !important;
    }}
    
    /* Custom Card Styling */
    .glass-card {{
        background: var(--card-bg) !important;
        border: 1px solid var(--border-color) !important;
        border-radius: 12px;
        padding: 1.25rem;
        box-shadow: 0 8px 32px 0 rgba(0, 0, 0, 0.05);
        margin-bottom: 1rem;
        transition: all 0.25s ease;
    }}
    
    .glass-card:hover {{
        border-color: rgba(16, 185, 129, 0.3) !important;
        box-shadow: 0 12px 40px 0 rgba(16, 185, 129, 0.05);
    }}
    
    /* Metrics Row */
    .metric-box {{
        display: flex;
        align-items: center;
        gap: 12px;
        background: var(--card-bg) !important;
        border: 1px solid var(--border-color) !important;
        border-radius: 10px;
        padding: 1rem 1.25rem;
        box-shadow: 0 4px 20px 0 rgba(0, 0, 0, 0.02);
    }}
    
    .metric-num {{
        font-size: 1.5rem;
        font-weight: 800;
        color: var(--title) !important;
    }}
    
    .metric-label {{
        font-size: 0.75rem;
        color: var(--text-muted) !important;
        text-transform: uppercase;
        letter-spacing: 0.05em;
    }}
    
    /* Custom Badges */
    .badge {{
        display: inline-block;
        padding: 2px 8px;
        border-radius: 9999px;
        font-size: 10px;
        font-weight: 700;
        text-transform: uppercase;
    }}
    .badge-completed {{ color: #10B981; background: rgba(16, 185, 129, 0.12); border: 1px solid rgba(16, 185, 129, 0.2); }}
    .badge-in-progress {{ color: #06B6D4; background: rgba(6, 182, 212, 0.12); border: 1px solid rgba(6, 182, 212, 0.2); }}
    .badge-pending {{ color: #F59E0B; background: rgba(245, 158, 11, 0.12); border: 1px solid rgba(245, 158, 11, 0.2); }}
    
    .badge-high {{ color: #EF4444; background: rgba(239, 68, 68, 0.1); border: 1px solid rgba(239, 68, 68, 0.2); }}
    .badge-medium {{ color: #F59E0B; background: rgba(245, 158, 11, 0.1); border: 1px solid rgba(245, 158, 11, 0.2); }}
    .badge-low {{ color: #475569; background: rgba(148, 163, 184, 0.1); border: 1px solid rgba(148, 163, 184, 0.25); }}

    /* Card Inner Typography */
    .card-title {{
        font-size: 13px;
        font-weight: 700;
        color: var(--title) !important;
        line-height: 1.4;
        margin-bottom: 6px;
    }}
    
    .card-title-completed {{
        font-size: 13px;
        font-weight: 700;
        color: var(--text-muted) !important;
        text-decoration: line-through;
        line-height: 1.4;
        margin-bottom: 6px;
    }}
    
    .card-desc {{
        font-size: 11px;
        color: var(--text-muted) !important;
        margin-bottom: 10px;
        line-height: 1.5;
    }}
    
    .card-date {{
        font-size: 10px;
        color: var(--text-muted) !important;
    }}
    
    .card-footer {{
        border-top: 1px solid var(--border-color) !important;
        padding-top: 8px;
        display: flex;
        justify-content: space-between;
        font-size: 11px;
        color: var(--text-muted) !important;
    }}
    
    .card-owner {{
        color: var(--title) !important;
        font-weight: 600;
    }}
    
    .card-meeting {{
        color: #06B6D4 !important;
        font-weight: 500;
    }}
    
    .card-progress {{
        background: var(--progress-bg) !important;
        border: 1px solid var(--progress-border) !important;
        padding: 6px 10px;
        border-radius: 6px;
        font-size: 11px;
        margin-top: 8px;
        color: #10B981 !important;
        line-height: 1.4;
    }}

    /* Pill tabs */
    div[data-baseweb="tab-list"] {{
        background: var(--card-bg) !important;
        border: 1px solid var(--border-color) !important;
        border-radius: 8px !important;
        padding: 3px !important;
    }}
    button[data-baseweb="tab"] {{
        color: var(--text-muted) !important;
        font-size: 0.8rem !important;
        font-weight: 600 !important;
        border-radius: 6px !important;
        border: none !important;
        padding: 6px 16px !important;
    }}
    button[data-baseweb="tab"][aria-selected="true"] {{
        color: #FFFFFF !important;
        background: #10B981 !important;
        box-shadow: 0 4px 12px rgba(16, 185, 129, 0.2) !important;
    }}
    [data-baseweb="tab-highlight"] {{
        display: none !important;
    }}
</style>
"""), unsafe_allow_html=True)

# 2. File paths and Local Storage Fallback
LOCAL_DATA_FILE = "tasks.json"
SETTINGS_FILE = "settings.json"

# Helper: Load Settings
def load_settings():
    if os.path.exists(SETTINGS_FILE):
        try:
            with open(SETTINGS_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except:
            pass
    return {"sheet_url": "", "gas_url": ""}

# Helper: Save Settings
def save_settings(settings):
    with open(SETTINGS_FILE, "w", encoding="utf-8") as f:
        json.dump(settings, f, ensure_ascii=False, indent=2)

# Helper: Parse Google Sheets export URL
def get_sheets_csv_url(url):
    if not url:
        return None
    match = re.search(r"/spreadsheets/d/([a-zA-Z0-9-_]+)", url)
    if match:
        spreadsheet_id = match.group(1)
        return f"https://docs.google.com/spreadsheets/d/{spreadsheet_id}/export?format=csv"
    return None

# Load Tasks & Migrate Schema dynamically
def load_tasks(settings):
    raw_tasks = []
    if settings.get("sheet_url"):
        csv_url = get_sheets_csv_url(settings["sheet_url"])
        if csv_url:
            try:
                df = pd.read_csv(csv_url)
                # Cast all values to string to prevent float NaN conversions
                df = df.astype(str).replace("nan", "")
                raw_tasks = df.to_dict(orient="records")
                # Write back as local cache for offline/resilience fallback
                save_tasks_safely(raw_tasks, LOCAL_DATA_FILE)
            except Exception as e:
                st.sidebar.error(f"⚠️ 無法讀取 Google 試算表，改為載入本地檔案。錯誤：{str(e)}")

    if not raw_tasks and os.path.exists(LOCAL_DATA_FILE):
        try:
            with open(LOCAL_DATA_FILE, "r", encoding="utf-8") as f:
                raw_tasks = json.load(f)
        except Exception as e:
            st.error(f"讀取 {LOCAL_DATA_FILE} 出錯：{str(e)}")

    # Migrate Schema dynamically (Self-healing data structure)
    migrated = False
    for t in raw_tasks:
        # Parse history strings from Google Sheets CSV
        if isinstance(t.get("meeting_history"), str):
            try:
                import json
                t["meeting_history"] = json.loads(t["meeting_history"])
            except Exception as e:
                pass
        if isinstance(t.get("progress_history"), str):
            try:
                import json
                t["progress_history"] = json.loads(t["progress_history"])
            except Exception as e:
                pass

        if "meeting_history" not in t or not isinstance(t["meeting_history"], list):
            t["meeting_history"] = [{
                "meeting": t.get("meeting", "策略發展會議"),
                "date": t.get("date", "2026.06.08"),
                "content": t.get("content", "")
            }]
            migrated = True
        if "progress_history" not in t or not isinstance(t["progress_history"], list):
            t["progress_history"] = [{
                "date": t.get("date", "2026.06.08"),
                "text": t.get("progress", "建立追蹤")
            }]
            migrated = True
        if "is_cross_dept" in t:
            val = t["is_cross_dept"]
            if isinstance(val, str):
                t["is_cross_dept"] = val.strip().lower() in ["true", "1", "yes", "t"]
                migrated = True
            elif isinstance(val, float):
                import math
                t["is_cross_dept"] = bool(val) and not math.isnan(val)
                migrated = True
            else:
                t["is_cross_dept"] = bool(val)
        else:
            is_cross = False
            for word in ["跨部門", "跨組", "跨單位", "協辦"]:
                if word in t.get("title", "") or word in t.get("content", ""):
                    is_cross = True
            t["is_cross_dept"] = is_cross
            migrated = True
    
    if migrated and raw_tasks:
        try:
            save_tasks_safely(raw_tasks, LOCAL_DATA_FILE)
        except:
            pass
            
    # Background retry queue: check if any tasks have sync_pending == True
    # And try to sync them now!
    gas_url = settings.get("gas_url")
    if gas_url and any(t.get("sync_pending") for t in raw_tasks):
        pending_tasks = [t for t in raw_tasks if t.get("sync_pending")]
        try:
            serializable_pending = []
            for t in pending_tasks:
                tc = t.copy()
                tc.pop("sync_pending", None)
                if "meeting_history" in tc and isinstance(tc["meeting_history"], list):
                    tc["meeting_history"] = json.dumps(tc["meeting_history"], ensure_ascii=False)
                if "progress_history" in tc and isinstance(tc["progress_history"], list):
                    tc["progress_history"] = json.dumps(tc["progress_history"], ensure_ascii=False)
                serializable_pending.append(tc)
                
            response = requests.post(
                gas_url,
                json=serializable_pending,
                headers={"Content-Type": "application/json"},
                timeout=5
            )
            if response.status_code == 200:
                for t in pending_tasks:
                    t["sync_pending"] = False
                save_tasks_safely(raw_tasks, LOCAL_DATA_FILE)
                st.sidebar.success("✅ 成功補齊先前漏接的進度！")
        except:
            pass

    return raw_tasks

# Save Tasks Safely (Atomic Write)
def save_tasks_safely(tasks, filepath):
    import tempfile
    fd, temp_path = tempfile.mkstemp(dir=os.path.dirname(os.path.abspath(filepath)))
    try:
        with os.fdopen(fd, 'w', encoding='utf-8') as f:
            json.dump(tasks, f, ensure_ascii=False, indent=2)
        os.replace(temp_path, filepath)
    except Exception as e:
        if os.path.exists(temp_path):
            os.remove(temp_path)
        raise e

# Save Tasks
def save_tasks(tasks, settings, modified_tasks=None):
    # 1. Always save raw tasks locally first (atomic save)
    save_tasks_safely(tasks, LOCAL_DATA_FILE)
    
    # 2. Check if cloud sync is configured
    gas_url = settings.get("gas_url")
    if not gas_url:
        return
        
    # If modified_tasks is [] (empty list), it means we only want to write locally (e.g. after deletion)
    if modified_tasks == []:
        return

    # Determine what to send (None = all tasks, otherwise modified_tasks)
    tasks_to_send = tasks if modified_tasks is None else modified_tasks
    
    # Mark them all as sync_pending = True in the actual tasks list
    for t in tasks_to_send:
        t["sync_pending"] = True
    
    # Save the pending state locally
    save_tasks_safely(tasks, LOCAL_DATA_FILE)
    
    # Serialize complex columns to JSON strings for Google Sheets
    serializable_tasks = []
    for t in tasks_to_send:
        tc = t.copy()
        tc.pop("sync_pending", None)
        if "meeting_history" in tc and isinstance(tc["meeting_history"], list):
            tc["meeting_history"] = json.dumps(tc["meeting_history"], ensure_ascii=False)
        if "progress_history" in tc and isinstance(tc["progress_history"], list):
            tc["progress_history"] = json.dumps(tc["progress_history"], ensure_ascii=False)
        serializable_tasks.append(tc)
        
    try:
        response = requests.post(
            gas_url,
            json=serializable_tasks,
            headers={"Content-Type": "application/json"},
            timeout=10
        )
        if response.status_code == 200:
            # Sync succeeded! Clear the sync_pending flag
            for t in tasks_to_send:
                t["sync_pending"] = False
            # Save the cleared state locally
            save_tasks_safely(tasks, LOCAL_DATA_FILE)
            st.sidebar.success("✅ 雲端 Google Sheets 同步成功！")
        else:
            st.sidebar.error("❌ 雲端同步失敗，已排入重試佇列。")
    except Exception as e:
        st.sidebar.error(f"⚠️ 網路超時或斷線，已排入重試佇列。")

# Initialize data
settings = load_settings()
if "tasks" not in st.session_state:
    st.session_state.tasks = load_tasks(settings)

if "limit_pending" not in st.session_state:
    st.session_state.limit_pending = 20
if "limit_progress" not in st.session_state:
    st.session_state.limit_progress = 20
if "limit_completed" not in st.session_state:
    st.session_state.limit_completed = 20

# Refresh helper
def refresh_data():
    st.session_state.tasks = load_tasks(settings)
    st.rerun()

# Helper: Get current logged-in user signature from Streamlit Cloud Context
def get_current_user_sig():
    try:
        if hasattr(st, "experimental_user"):
            u = st.experimental_user
            email = getattr(u, "email", None)
            if not email and hasattr(u, "get"):
                email = u.get("email")
            if email:
                return email.split("@")[0]
    except Exception:
        pass
    return "本地管理員"


# 3. Header Panel
head_left, head_right = st.columns([9, 1])
with head_left:
    st.markdown(clean_html(f"""
    <div style="display: flex; align-items: center; gap: 15px; margin-bottom: 1.5rem;">
        <div style="background: #10B981; color: white; width: 38px; height: 38px; border-radius: 10px; display: flex; align-items: center; justify-content: center; font-weight: bold; font-size: 20px; box-shadow: 0 4px 14px rgba(16, 185, 129, 0.4);">
            P
        </div>
        <div>
            <h1 style="font-size: 1.5rem; font-weight: 800; margin: 0; color: {TITLE_COLOR}; letter-spacing: -0.02em;">博幼會議專案待辦與進度追蹤</h1>
            <p style="font-size: 0.75rem; margin: 0; color: #10B981; font-weight: 700; letter-spacing: 0.05em; text-transform: uppercase;">教育輔導處行政幕僚與專案管理系統</p>
        </div>
    </div>
    """), unsafe_allow_html=True)

with head_right:
    theme_btn = "☀️ 白天模式" if IS_DARK else "🌙 暗黑模式"
    if st.button(theme_btn, use_container_width=True):
        st.session_state.theme = "light" if IS_DARK else "dark"
        st.rerun()

# 4. Metrics Dashboard Row
tasks = st.session_state.tasks
total_cnt = len(tasks)
pending_cnt = len([t for t in tasks if t.get("status") == "pending"])
progress_cnt = len([t for t in tasks if t.get("status") == "in_progress"])
completed_cnt = len([t for t in tasks if t.get("status") == "completed"])
completion_rate = round((completed_cnt / total_cnt * 100)) if total_cnt > 0 else 0

m1, m2, m3, m4, m5 = st.columns(5)
with m1:
    st.markdown(clean_html(f'<div class="metric-box"><div style="background: rgba(99,102,241,0.15); color: #6366F1; padding: 8px; border-radius: 8px; font-size: 18px;">📋</div><div><div class="metric-label">任務總數</div><div class="metric-num">{total_cnt}</div></div></div>'), unsafe_allow_html=True)
with m2:
    st.markdown(clean_html(f'<div class="metric-box"><div style="background: rgba(245,158,11,0.15); color: #F59E0B; padding: 8px; border-radius: 8px; font-size: 18px;">⏳</div><div><div class="metric-label">待處理</div><div class="metric-num">{pending_cnt}</div></div></div>'), unsafe_allow_html=True)
with m3:
    st.markdown(clean_html(f'<div class="metric-box"><div style="background: rgba(6,182,212,0.15); color: #06B6D4; padding: 8px; border-radius: 8px; font-size: 18px;">⚡</div><div><div class="metric-label">進行中</div><div class="metric-num">{progress_cnt}</div></div></div>'), unsafe_allow_html=True)
with m4:
    st.markdown(clean_html(f'<div class="metric-box"><div style="background: rgba(16,185,129,0.15); color: #10B981; padding: 8px; border-radius: 8px; font-size: 18px;">✅</div><div><div class="metric-label">已完成</div><div class="metric-num">{completed_cnt}</div></div></div>'), unsafe_allow_html=True)
with m5:
    st.markdown(clean_html(f'<div class="metric-box" style="flex-direction: column; align-items: stretch; justify-content: center; gap: 4px;"><div style="display: flex; justify-content: space-between; font-size: 11px;"><span class="metric-label">任務完成率</span><span style="color:#10B981; font-weight:700;">{completion_rate}%</span></div><div style="background: #E2E8F0; border-radius: 9999px; height: 6px; overflow: hidden;"><div style="background: #10B981; width: {completion_rate}%; height: 100%; border-radius: 9999px;"></div></div></div>'), unsafe_allow_html=True)

st.markdown("<div style='margin-bottom: 1.5rem;'></div>", unsafe_allow_html=True)

# 5. Sidebar Filters
st.sidebar.markdown("### 🔍 任務篩選與控制")
search_query = st.sidebar.text_input("搜尋工作名稱、負責人或會議內容...")
filter_priority = st.sidebar.selectbox("任務優先級", ["All", "high", "medium", "low"])
filter_status = st.sidebar.selectbox("任務狀態", ["All", "pending", "in_progress", "completed", "archived"])

st.sidebar.markdown("---")
if st.sidebar.button("🔄 重新載入 / 同步資料庫", use_container_width=True):
    refresh_data()

# Apply Sidebar filters to data
filtered_tasks = tasks
if search_query:
    filtered_tasks = [t for t in filtered_tasks if search_query.lower() in t.get("title", "").lower() or search_query.lower() in t.get("owner", "").lower() or search_query.lower() in t.get("content", "").lower()]
if filter_priority != "All":
    filtered_tasks = [t for t in filtered_tasks if t.get("priority") == filter_priority]
if filter_status != "All":
    filtered_tasks = [t for t in filtered_tasks if t.get("status") == filter_status]

# 6. Main Navigation Tabs (Added 📦 歷史歸檔檔案庫 Tab)
tab_kanban, tab_parser, tab_dictionary, tab_archive, tab_sync = st.tabs([
    "📋 任務追蹤看板", 
    "📥 會議紀錄匯入", 
    "📖 業務分類字典與協作指南", 
    "📦 歷史歸檔檔案庫",
    "⚙️ 雲端同步設定"
])

# ---- DIALOG: EDIT TASK ----
if hasattr(st, "dialog"):
    @st.dialog("✏️ 編輯與更新任務進度")
    def show_edit_dialog(task):
        st.write(f"📌 正在更新項目：**{task.get('title')}**")
        st.markdown("---")
        
        ut_title = st.text_input("任務名稱", value=task.get("title"))
        ut_dept = st.selectbox("主責部門", ["教材研發組", "社工特教組", "學區營運組", "畢業生組", "處長室/行政管理"], index=["教材研發組", "社工特教組", "學區營運組", "畢業生組", "處長室/行政管理"].index(task.get("dept")))
        ut_status = st.selectbox("進度狀態", ["pending", "in_progress", "completed", "archived"], index=["pending", "in_progress", "completed", "archived"].index(task.get("status")), format_func=lambda x: "⏳ 待處理" if x == "pending" else ("⚡ 進行中" if x == "in_progress" else ("✅ 已完成" if x == "completed" else "📦 已歸檔")))
        ut_owner = st.text_input("主責負責人", value=task.get("owner"))
        ut_priority = st.selectbox("優先度", ["high", "medium", "low"], index=["high", "medium", "low"].index(task.get("priority")))
        ut_date = st.text_input("最近更新日期", value=pd.Timestamp.now().strftime('%Y.%m.%d'))
        ut_content = st.text_area("詳細決議 / 任務說明", value=task.get("content"))
        ut_cross = st.checkbox("🤝 此任務需要跨部門/跨組別協作配合", value=bool(task.get("is_cross_dept", False)))
        
        # Automatic user signature extraction
        current_user = get_current_user_sig()
        st.info(f"👤 目前操作帳號：**{current_user}** (系統將自動關聯修改紀錄，無須手動簽名)")
        
        st.markdown("---")
        st.subheader("📈 新增執行進度")
        new_progress = st.text_input("填寫此任務的最新執行進度：", placeholder="填寫後點選儲存會寫入歷史進度中...")
        
        st.markdown("---")
        confirm_save = st.checkbox("⚠️ 我已確認上述修改與進度無誤，同意儲存與同步", value=False)
        
        col_c1, col_c2 = st.columns(2)
        with col_c1:
            if st.button("💾 儲存修改進度", type="primary", disabled=not confirm_save, use_container_width=True):
                # Track modified fields for audit log
                changes = []
                if task.get("title") != ut_title: changes.append("修改名稱")
                if task.get("dept") != ut_dept: changes.append("修改部門")
                if task.get("owner") != ut_owner: changes.append(f"負責人改為 {ut_owner}")
                if task.get("status") != ut_status:
                    status_map = {"pending": "待處理", "in_progress": "進行中", "completed": "已完成", "archived": "已歸檔"}
                    changes.append(f"狀態改為 {status_map.get(ut_status, ut_status)}")
                if task.get("priority") != ut_priority: changes.append(f"調整優先度為 {ut_priority}")
                if bool(task.get("is_cross_dept")) != bool(ut_cross): changes.append("調整跨部門設定")

                task["title"] = ut_title
                task["dept"] = ut_dept
                task["owner"] = ut_owner
                task["status"] = ut_status
                task["priority"] = ut_priority
                task["date"] = ut_date
                task["content"] = ut_content
                task["is_cross_dept"] = ut_cross
                
                updater_sig = current_user
                if new_progress.strip():
                    full_txt = f"[手動更新 - {updater_sig}] {new_progress.strip()}"
                    task["progress"] = full_txt
                    if "progress_history" not in task:
                        task["progress_history"] = []
                    task["progress_history"].insert(0, {
                        "date": ut_date,
                        "text": full_txt
                    })
                elif changes:
                    changes_str = "、".join(changes)
                    full_txt = f"[屬性異動 - {updater_sig}] {changes_str}"
                    task["progress"] = full_txt
                    if "progress_history" not in task:
                        task["progress_history"] = []
                    task["progress_history"].insert(0, {
                        "date": ut_date,
                        "text": full_txt
                    })
                
                for idx, t in enumerate(st.session_state.tasks):
                    if t["id"] == task["id"]:
                        st.session_state.tasks[idx] = task
                        break
                
                save_tasks(st.session_state.tasks, settings, modified_tasks=[task])
                st.success("✅ 儲存成功！")
                st.rerun()
                
        with col_c2:
            if st.button("🗑️ 刪除此項任務", type="secondary", use_container_width=True):
                task["status"] = "deleted"
                save_tasks(st.session_state.tasks, settings, modified_tasks=[task])
                st.session_state.tasks = [t for t in st.session_state.tasks if t["id"] != task["id"]]
                save_tasks(st.session_state.tasks, settings, modified_tasks=[])
                st.warning("⚠️ 任務已從追蹤清單中刪除。")
                st.rerun()

        # Feature 1: Progress history view and one-click rollback mechanism
        st.markdown("---")
        st.markdown("#### ↩️ 歷史進度一鍵復原庫")
        ph_list = task.get("progress_history", [])
        if ph_list:
            for p_idx, ph in enumerate(ph_list):
                col_ph_text, col_ph_btn = st.columns([7.5, 2.5])
                with col_ph_text:
                    st.markdown(clean_html(f"""
                    <div style='font-size: 11px; line-height: 1.45; color: var(--text-muted); margin-bottom:6px;'>
                        <b>{html_escape(ph.get('date'))}</b>: {html_escape(ph.get('text'))}
                    </div>
                    """), unsafe_allow_html=True)
                with col_ph_btn:
                    if st.button("↩️ 復原此時", key=f"rb_{task['id']}_{p_idx}", help="將工作進度倒轉/復原至此時狀態", use_container_width=True):
                        raw_txt = ph.get("text", "")
                        clean_txt = re.sub(r"^\[.*?\]\s*", "", raw_txt)
                        rollback_txt = f"[復原進度] {clean_txt}"
                        
                        task["progress"] = rollback_txt
                        task["date"] = pd.Timestamp.now().strftime('%Y.%m.%d')
                        if "progress_history" not in task:
                            task["progress_history"] = []
                        task["progress_history"].insert(0, {
                            "date": task["date"],
                            "text": rollback_txt
                        })
                        
                        for idx, t in enumerate(st.session_state.tasks):
                            if t["id"] == task["id"]:
                                st.session_state.tasks[idx] = task
                                break
                        save_tasks(st.session_state.tasks, settings, modified_tasks=[task])
                        st.success("✅ 已復原進度狀態！")
                        st.rerun()
        else:
            st.info("尚無歷史進度紀錄。")
else:
    def show_edit_dialog(task):
        st.sidebar.markdown("---")
        st.sidebar.markdown(f"### ✏️ 編輯任務: {task.get('title')}")
        ut_title = st.sidebar.text_input("任務名稱", value=task.get("title"))
        ut_owner = st.sidebar.text_input("主責對口", value=task.get("owner"))
        ut_status = st.sidebar.selectbox("進度狀態", ["pending", "in_progress", "completed", "archived"], index=["pending", "in_progress", "completed", "archived"].index(task.get("status")))
        ut_cross = st.sidebar.checkbox("跨部門協作任務", value=bool(task.get("is_cross_dept", False)))
        current_user = get_current_user_sig()
        st.sidebar.info(f"👤 目前操作帳號：{current_user}")
        new_progress = st.sidebar.text_input("最新進度說明", value="")
        confirm_save = st.sidebar.checkbox("確認儲存", value=False)
        
        if st.sidebar.button("確認保存", disabled=not confirm_save):
            # Track modified fields for audit log
            changes = []
            if task.get("title") != ut_title: changes.append("修改名稱")
            if task.get("owner") != ut_owner: changes.append(f"負責人改為 {ut_owner}")
            if task.get("status") != ut_status:
                status_map = {"pending": "待處理", "in_progress": "進行中", "completed": "已完成", "archived": "已歸檔"}
                changes.append(f"狀態改為 {status_map.get(ut_status, ut_status)}")
            if bool(task.get("is_cross_dept")) != bool(ut_cross): changes.append("調整跨部門設定")

            task["title"] = ut_title
            task["owner"] = ut_owner
            task["status"] = ut_status
            task["is_cross_dept"] = ut_cross
            
            updater_sig = current_user
            cur_date = pd.Timestamp.now().strftime("%Y.%m.%d")
            if new_progress.strip():
                full_txt = f"[手動更新 - {updater_sig}] {new_progress.strip()}"
                task["progress"] = full_txt
                task["progress_history"].insert(0, {"date": cur_date, "text": full_txt})
            elif changes:
                changes_str = "、".join(changes)
                full_txt = f"[屬性異動 - {updater_sig}] {changes_str}"
                task["progress"] = full_txt
                task["progress_history"].insert(0, {"date": cur_date, "text": full_txt})
                
            for idx, t in enumerate(st.session_state.tasks):
                if t["id"] == task["id"]:
                    st.session_state.tasks[idx] = task
                    break
            save_tasks(st.session_state.tasks, settings, modified_tasks=[task])
            st.rerun()

# ---- DIALOG: CREATE TASK ----
if hasattr(st, "dialog"):
    @st.dialog("➕ 新增全新專案任務")
    def show_create_dialog():
        st.markdown("#### 🆕 填寫新任務內容")
        nt_title = st.text_input("任務名稱 *", placeholder="請輸入任務名稱")
        nt_dept = st.selectbox("主責部門", ["教材研發組", "社工特教組", "學區營運組", "畢業生組", "處長室/行政管理"])
        nt_owner = st.text_input("主責對口 / 負責人", placeholder="例如：韋志 / 睿豪")
        nt_status = st.selectbox("進度狀態", ["pending", "in_progress", "completed"], format_func=lambda x: "⏳ 待處理" if x == "pending" else ("⚡ 進行中" if x == "in_progress" else "✅ 已完成"))
        nt_priority = st.selectbox("優先度", ["high", "medium", "low"])
        nt_date = st.text_input("會議日期 / 期限", value=pd.Timestamp.now().strftime('%Y.%m.%d'))
        nt_content = st.text_area("詳細決議 / 任務說明", placeholder="請輸入此任務的詳細內容與要求細節...")
        nt_progress = st.text_input("初始進度說明", value="建立追蹤")
        nt_cross = st.checkbox("🤝 此任務需要跨部門/跨組別協作配合", value=False)
        current_user = get_current_user_sig()
        st.info(f"👤 目前操作帳號：**{current_user}** (系統將自動關聯建立紀錄，無須手動簽名)")
        
        st.markdown("---")
        confirm_add = st.checkbox("我已確認填寫內容無誤", value=False)
        
        if st.button("➕ 確認新增待辦項目", type="primary", disabled=not confirm_add, use_container_width=True):
            if not nt_title:
                st.error("請輸入任務名稱！")
            else:
                creator_sig = current_user
                initial_prog = f"[手動新增 - {creator_sig}] {nt_progress.strip()}"
                new_item = {
                    "id": f"task-manual-{int(pd.Timestamp.now().timestamp())}",
                    "title": nt_title,
                    "dept": nt_dept,
                    "owner": nt_owner or "待指派",
                    "meeting": "行政手動新增",
                    "date": nt_date,
                    "status": nt_status,
                    "priority": nt_priority,
                    "content": nt_content or nt_title,
                    "progress": initial_prog,
                    "is_cross_dept": nt_cross,
                    "created_date": nt_date,
                    "meeting_history": [{
                        "meeting": "行政手動新增",
                        "date": nt_date,
                        "content": nt_content or nt_title
                    }],
                    "progress_history": [{
                        "date": nt_date,
                        "text": initial_prog
                    }]
                }
                st.session_state.tasks.insert(0, new_item)
                save_tasks(st.session_state.tasks, settings, modified_tasks=[new_item])
                st.success(f"✅ 成功手動新增任務：{nt_title}")
                st.rerun()

# Setup editing session trigger
if "active_edit_task" not in st.session_state:
    st.session_state.active_edit_task = None

if st.session_state.active_edit_task:
    task_to_edit = st.session_state.active_edit_task
    st.session_state.active_edit_task = None
    show_edit_dialog(task_to_edit)

# ---- TAB 1: 看板 ----
with tab_kanban:
    col_k_head, col_k_layout, col_k_btn = st.columns([6, 3, 3])
    with col_k_head:
        selected_subgroup = st.radio("主責業務部門分組：", ["全部", "教材研發組", "社工特教組", "學區營運組", "畢業生組", "處長室/行政管理"], horizontal=True)
    with col_k_layout:
        layout_mode = st.radio("選擇呈現版面：", ["📋 任務看板 (Kanban)", "📊 全覽表格 (Table)"], horizontal=True)
    with col_k_btn:
        if st.button("➕ 新增全新專案任務", type="primary", use_container_width=True):
            if hasattr(st, "dialog"):
                show_create_dialog()
            else:
                st.info("請到側邊欄填寫新增任務")

    # Exclude archived tasks and completed tasks older than 3 months from main Kanban board and active overview table
    import pandas as pd
    kanban_tasks = []
    for t in filtered_tasks:
        if t.get("status") == "archived":
            continue
        if t.get("status") == "completed":
            try:
                task_date = pd.to_datetime(t.get("date", ""), format="%Y.%m.%d")
                three_months_ago = pd.Timestamp.now() - pd.DateOffset(months=3)
                if task_date < three_months_ago:
                    continue
            except:
                pass
        kanban_tasks.append(t)
    if selected_subgroup != "全部":
        kanban_tasks = [t for t in kanban_tasks if t.get("dept") == selected_subgroup]

    # CASE A: Render Kanban Layout
    if layout_mode == "📋 任務看板 (Kanban)":
        col_pending, col_progress, col_completed = st.columns(3)
        
        def render_kanban_card(t, index_key):
            # Escape strings to prevent HTML/XSS leakage
            esc_title = html_escape(t.get("title", ""))
            esc_content = html_escape(t.get("content", ""))
            esc_progress = html_escape(t.get("progress", ""))
            esc_owner = html_escape(t.get("owner", ""))
            esc_date = html_escape(t.get("date", ""))
            esc_dept = html_escape(t.get("dept", ""))
            esc_priority = html_escape(t.get("priority", ""))

            # Calculate project duration
            start_date = get_task_start_date(t)
            if t.get("status") in ["completed", "archived"]:
                period_str = f"{start_date} ~ {esc_date} (已完結)"
            else:
                period_str = f"{start_date} ~ 至今"

            with st.container(border=True):
                c_badges, c_edit = st.columns([10, 2])
                with c_badges:
                    pri_badge = f'<span class="badge badge-{esc_priority}">{esc_priority}</span>'
                    dept_badge = f'<span class="badge" style="background: rgba(99,102,241,0.1); color: #6366F1; border: 1px solid rgba(99,102,241,0.15); margin-left: 5px;">{esc_dept}</span>'
                    cross_badge = '<span class="badge" style="background: rgba(244,63,94,0.1); color: #F43F5E; border: 1px solid rgba(244,63,94,0.15); margin-left: 5px;">🤝 跨部門</span>' if t.get("is_cross_dept") else ''
                    st.markdown(clean_html(f"{pri_badge}{dept_badge}{cross_badge}"), unsafe_allow_html=True)
                with c_edit:
                    if st.button("✏️", key=f"edit_btn_{t['id']}_{index_key}", help="編輯此任務與更新進度", use_container_width=True):
                        st.session_state.active_edit_task = t
                        st.rerun()

                # Build collapsible history HTML blocks
                mh_list = t.get("meeting_history", [])
                meeting_history_html = "".join([f'<li style="margin-bottom:4px;"><b>{html_escape(mh.get("date"))}</b>: {html_escape(mh.get("meeting"))} - <i>{html_escape(mh.get("content"))}</i></li>' for mh in mh_list])
                
                ph_list = t.get("progress_history", [])
                progress_history_html = "".join([f'<div style="margin-bottom:4px; padding-bottom:4px; border-bottom:1px dashed var(--border-color);"><b>{html_escape(ph.get("date"))}</b>: {html_escape(ph.get("text"))}</div>' for ph in ph_list])

                cross_suggestion_html = ""

                title_class = "card-title-completed" if t.get("status") == "completed" else "card-title"

                st.markdown(clean_html(f"""
                <div class="{title_class}" style="margin-top: 6px;">{esc_title}</div>
                <div class="card-desc">{esc_content}</div>
                
                <div style="font-size: 11px; color: var(--text-muted); margin-top: 4px; margin-bottom: 6px;">
                    📅 <b>專案期程：</b>{period_str}
                </div>
                
                <div class="card-progress">
                    <b>最新進度：</b>{esc_progress}
                </div>
                
                
                <div style="margin-top: 8px; border-top: 1px solid var(--border-color); padding-top: 8px;">
                    <details style="cursor: pointer; margin-bottom: 6px;">
                        <summary style="font-size: 11px; color: #10B981; font-weight:600;">📈 歷史執行進度追蹤 ({len(ph_list)}條)</summary>
                        <div style="margin: 5px 0 0 5px; padding-left: 8px; border-left: 2px solid #10B981; font-size: 11px; color: var(--text-muted);">
                            {progress_history_html}
                        </div>
                    </details>
                </div>
                
                <div>
                    <details style="cursor: pointer;">
                        <summary style="font-size: 11px; color: #06B6D4; font-weight:600;">📂 歷史提及會議紀錄 ({len(mh_list)}次)</summary>
                        <ul style="margin: 5px 0 0 10px; padding-left: 10px; font-size: 11px; color: var(--text-muted);">
                            {meeting_history_html}
                        </ul>
                    </details>
                </div>
                
                <div class="card-footer" style="margin-top: 8px; margin-bottom: 8px;">
                    <span>主責：<b class="card-owner">{esc_owner}</b></span>
                    <span class="card-date">最後更新：{esc_date}</span>
                </div>
                """), unsafe_allow_html=True)

                col_btn1, col_btn2 = st.columns(2)
                with col_btn1:
                    if t.get("status") == "pending":
                        if st.button("⚡ 開始執行", key=f"start_{t['id']}_{index_key}", use_container_width=True):
                            t["status"] = "in_progress"
                            t["progress"] = "狀態更新：開始執行項目"
                            t["date"] = pd.Timestamp.now().strftime("%Y.%m.%d")
                            t["progress_history"].insert(0, {
                                "date": t["date"],
                                "text": "狀態更新：開始執行項目"
                            })
                            save_tasks(st.session_state.tasks, settings, modified_tasks=[t])
                            st.rerun()
                    elif t.get("status") == "in_progress":
                        if st.button("⏳ 移回待辦", key=f"revert_{t['id']}_{index_key}", use_container_width=True):
                            t["status"] = "pending"
                            t["progress"] = "狀態更新：移回待辦狀態"
                            t["date"] = pd.Timestamp.now().strftime("%Y.%m.%d")
                            t["progress_history"].insert(0, {
                                "date": t["date"],
                                "text": "狀態更新：移回待辦狀態"
                            })
                            save_tasks(st.session_state.tasks, settings, modified_tasks=[t])
                            st.rerun()
                    elif t.get("status") == "completed":
                        if st.button("⚡ 重啟任務", key=f"reopen_{t['id']}_{index_key}", use_container_width=True):
                            t["status"] = "in_progress"
                            t["progress"] = "狀態更新：重啟任務項目"
                            t["date"] = pd.Timestamp.now().strftime("%Y.%m.%d")
                            t["progress_history"].insert(0, {
                                "date": t["date"],
                                "text": "狀態更新：重啟任務項目"
                            })
                            save_tasks(st.session_state.tasks, settings, modified_tasks=[t])
                            st.rerun()
                with col_btn2:
                    if t.get("status") in ["pending", "in_progress"]:
                        if st.button("✅ 標記完成", key=f"complete_{t['id']}_{index_key}", use_container_width=True):
                            t["status"] = "completed"
                            t["progress"] = "狀態更新：標記完成"
                            t["date"] = pd.Timestamp.now().strftime("%Y.%m.%d")
                            t["progress_history"].insert(0, {
                                "date": t["date"],
                                "text": "狀態更新：標記完成"
                            })
                            save_tasks(st.session_state.tasks, settings, modified_tasks=[t])
                            st.rerun()
                    elif t.get("status") == "completed":
                        if st.button("📦 歸檔任務", key=f"archive_{t['id']}_{index_key}", help="歸檔後移出主看板，可在歷史檔案庫查閱", use_container_width=True):
                            t["status"] = "archived"
                            t["progress"] = "狀態更新：已手動歸檔此任務"
                            t["date"] = pd.Timestamp.now().strftime("%Y.%m.%d")
                            t["progress_history"].insert(0, {
                                "date": t["date"],
                                "text": "狀態更新：已手動歸檔此任務"
                            })
                            save_tasks(st.session_state.tasks, settings, modified_tasks=[t])
                            st.rerun()

        with col_pending:
            st.markdown("<h4 style='color:#F59E0B; border-bottom: 2px solid #F59E0B; padding-bottom: 6px; font-weight:700;'>&#9203; &#24453;&#34389;&#29702;</h4>", unsafe_allow_html=True)
            pending_list = [t for t in kanban_tasks if t.get("status") == "pending"]
            display_pending = pending_list[:st.session_state.limit_pending]
            for idx, t in enumerate(display_pending):
                render_kanban_card(t, f"pend_{idx}")
            if not pending_list:
                st.info("&#28961;&#27492;&#39006;&#21029;&#24453;&#36774;&#38917;&#30446;")
            elif len(pending_list) > st.session_state.limit_pending:
                if st.button("載入更多待辦項目...", key="btn_load_more_pending", use_container_width=True):
                    st.session_state.limit_pending += 20
                    st.rerun()

        with col_progress:
            st.markdown("<h4 style='color:#06B6D4; border-bottom: 2px solid #06B6D4; padding-bottom: 6px; font-weight:700;'>&#9889; &#36914;&#34892;&#20013;</h4>", unsafe_allow_html=True)
            progress_list = [t for t in kanban_tasks if t.get("status") == "in_progress"]
            display_progress = progress_list[:st.session_state.limit_progress]
            for idx, t in enumerate(display_progress):
                render_kanban_card(t, f"prog_{idx}")
            if not progress_list:
                st.info("&#28961;&#36914;&#34892;&#20013;&#38917;&#30446;")
            elif len(progress_list) > st.session_state.limit_progress:
                if st.button("載入更多進行中項目...", key="btn_load_more_progress", use_container_width=True):
                    st.session_state.limit_progress += 20
                    st.rerun()

        with col_completed:
            st.markdown("<h4 style='color:#10B981; border-bottom: 2px solid #10B981; padding-bottom: 6px; font-weight:700;'>&#9989; &#24050;&#23436;&#25104;</h4>", unsafe_allow_html=True)
            completed_list = [t for t in kanban_tasks if t.get("status") == "completed"]
            display_completed = completed_list[:st.session_state.limit_completed]
            for idx, t in enumerate(display_completed):
                render_kanban_card(t, f"comp_{idx}")
            if not completed_list:
                st.info("&#28961;&#24050;&#23436;&#25104;&#38917;&#30446;")
            elif len(completed_list) > st.session_state.limit_completed:
                if st.button("載入更多已完成項目...", key="btn_load_more_completed", use_container_width=True):
                    st.session_state.limit_completed += 20
                    st.rerun()

    # CASE B: Render Full Overview Table Layout
    else:
        st.markdown("<div style='margin-bottom: 1rem;'></div>", unsafe_allow_html=True)
        
        # Grid Header Row
        t_cols = st.columns([1.2, 1.5, 3.2, 1.2, 0.8, 3.1, 1.0])
        t_cols[0].markdown("**狀態**")
        t_cols[1].markdown("**主責部門**")
        t_cols[2].markdown("**任務項目 / 決議說明**")
        t_cols[3].markdown("**負責人**")
        t_cols[4].markdown("**協作**")
        t_cols[5].markdown("**最新執行進度與說明**")
        t_cols[6].markdown("**操作**")
        st.markdown("<hr style='margin: 0.5rem 0 1rem 0; border-color: var(--border-color);'>", unsafe_allow_html=True)
        
        # Grid Data Rows
        for idx, t in enumerate(kanban_tasks):
            # Escape strings to prevent layout breakdown
            esc_title = html_escape(t.get("title", ""))
            esc_content = html_escape(t.get("content", ""))
            esc_progress = html_escape(t.get("progress", ""))
            esc_owner = html_escape(t.get("owner", ""))
            esc_dept = html_escape(t.get("dept", ""))
            esc_date = html_escape(t.get("date", ""))

            # Calculate project duration
            start_date = get_task_start_date(t)
            if t.get("status") in ["completed", "archived"]:
                period_str = f"{start_date} ~ {esc_date} (已完結)"
            else:
                period_str = f"{start_date} ~ 至今"

            r_cols = st.columns([1.2, 1.5, 3.2, 1.2, 0.8, 3.1, 1.0])
            
            # Column 1: Status badge
            status_html = ""
            if t.get("status") == "pending":
                status_html = '<span class="badge badge-pending">⏳ 待處理</span>'
            elif t.get("status") == "in_progress":
                status_html = '<span class="badge badge-in-progress">⚡ 進行中</span>'
            elif t.get("status") == "completed":
                status_html = '<span class="badge badge-completed">✅ 已完成</span>'
            r_cols[0].markdown(clean_html(status_html), unsafe_allow_html=True)
            
            # Column 2: Department badge
            dept_html = f'<span class="badge" style="background: rgba(99,102,241,0.08); color: #6366F1; border: 1px solid rgba(99,102,241,0.15);">{esc_dept}</span>'
            r_cols[1].markdown(clean_html(dept_html), unsafe_allow_html=True)
            
            # Column 3: Task & description & meeting history dropdown
            mh_list = t.get("meeting_history", [])
            meeting_history_html = "".join([f'<li style="margin-bottom:3px;"><b>{html_escape(mh.get("date"))}</b>: {html_escape(mh.get("meeting"))} - <i>{html_escape(mh.get("content"))}</i></li>' for mh in mh_list])
            title_style = "text-decoration: line-through; color: var(--text-muted);" if t.get("status") == "completed" else "color: var(--title);"
            
            r_cols[2].markdown(clean_html(f"""
            <div style="font-weight: 700; font-size: 13px; {title_style}">{esc_title}</div>
            <div style="font-size: 11px; color: var(--text-muted); margin-top: 2px; line-height:1.4;">{esc_content}</div>
            <div style="font-size: 10.5px; color: var(--text-muted); margin-top: 3px; font-weight: 500;">📅 期程: {period_str}</div>
            <details style="cursor: pointer; margin-top: 4px;">
                <summary style="font-size: 10px; color: #06B6D4; font-weight:600;">📂 提及歷史 ({len(mh_list)}次)</summary>
                <ul style="margin: 4px 0 0 10px; padding-left: 10px; font-size: 10.5px; color: var(--text-muted); line-height:1.45;">
                    {meeting_history_html}
                </ul>
            </details>
            """), unsafe_allow_html=True)
            
            # Column 4: Owner
            r_cols[3].markdown(clean_html(f"<div style='font-size: 12px; font-weight:600; color: var(--title); padding-top:4px;'>{esc_owner}</div>"), unsafe_allow_html=True)
            
            # Column 5: Cross-dept badge
            cross_html = "🤝 跨部門" if t.get("is_cross_dept") else "-"
            cross_style = "color: #F43F5E; font-weight:700; font-size:11px;" if t.get("is_cross_dept") else "color: var(--text-muted); font-size:11px;"
            r_cols[4].markdown(clean_html(f"<div style='{cross_style}; padding-top:4px;'>{cross_html}</div>"), unsafe_allow_html=True)
            
            # Column 6: Progress & progress history dropdown
            ph_list = t.get("progress_history", [])
            progress_history_html = "".join([f'<div style="margin-bottom:4px; padding-bottom:4px; border-bottom:1px dashed var(--border-color);"><b>{html_escape(ph.get("date"))}</b>: {html_escape(ph.get("text"))}</div>' for ph in ph_list])
            
            r_cols[5].markdown(clean_html(f"""
            <div style="font-size: 12px; color: #10B981; font-weight:600; padding-top:4px;">{esc_progress}</div>
            <details style="cursor: pointer; margin-top: 4px;">
                <summary style="font-size: 10px; color: #10B981; font-weight:600;">📈 歷史進度 ({len(ph_list)}條)</summary>
                <div style="margin: 4px 0 0 0; padding-left: 6px; border-left: 2px solid #10B981; font-size: 10.5px; color: var(--text-muted); line-height:1.45;">
                    {progress_history_html}
                </div>
            </details>
            """), unsafe_allow_html=True)
            
            # Column 7: Action (Edit & Archive side-by-side)
            with r_cols[6]:
                tbl_btn_c1, tbl_btn_c2 = st.columns([1, 1])
                with tbl_btn_c1:
                    if st.button("✏️", key=f"tbl_edit_{t['id']}_{idx}", help="編輯此任務", use_container_width=True):
                        st.session_state.active_edit_task = t
                        st.rerun()
                with tbl_btn_c2:
                    if t.get("status") == "completed":
                        if st.button("📦", key=f"tbl_arc_{t['id']}_{idx}", help="歸檔此任務", use_container_width=True):
                            t["status"] = "archived"
                            t["progress"] = "狀態更新：已手動歸檔此任務"
                            t["date"] = pd.Timestamp.now().strftime("%Y.%m.%d")
                            t["progress_history"].insert(0, {
                                "date": t["date"],
                                "text": "狀態更新：已手動歸檔此任務"
                            })
                            save_tasks(st.session_state.tasks, settings, modified_tasks=[t])
                            st.rerun()
            
            # Divider between rows
            st.markdown("<hr style='margin: 0.5rem 0; border-color: var(--border-color); opacity: 0.5;'>", unsafe_allow_html=True)

# ---- TAB 2: 會議紀錄匯入 ----
with tab_parser:
    st.markdown("### 📥 自動提取新會議待辦")
    st.markdown("上傳 Word 檔 (`.docx`) 或貼上會議記錄文字，系統將自動解析出待辦事項，並提供篩選與匯入。")
    
    parser_mode = st.radio("選擇會議導入方式：", ["上傳 Word 檔案 (.docx)", "直接貼上會議紀錄文字"], horizontal=True)

    extracted_tasks = []

    if parser_mode == "上傳 Word 檔案 (.docx)":
        uploaded_file = st.file_uploader("請選擇 Word 會議紀錄檔案", type="docx")
        if uploaded_file:
            with st.spinner("正在提取會議紀錄中的待辦表格與事項..."):
                file_bytes = uploaded_file.read()
                extracted_tasks = parse_docx_to_tasks(file_bytes, uploaded_file.name)
    else:
        text_input = st.text_area("請在此貼上您的會議紀錄內容...", height=200, placeholder="例如：\n教材研發組 | 進行國二教材勘誤 | 進行中\n社工特教組 | 調整兒少保護流程 | 6/26分享完成")
        if st.button("開始解析文字", type="primary"):
            lines = text_input.split("\n")
            for idx, line in enumerate(lines):
                trimmed = line.strip()
                if len(trimmed) > 8:
                    detected_dept = detect_department(trimmed)
                    extracted_tasks.append({
                        "id": f"parsed-text-{int(pd.Timestamp.now().timestamp())}-{idx}",
                        "title": trimmed[:60] + "..." if len(trimmed) > 60 else trimmed,
                        "dept": detected_dept,
                        "owner": "待指派",
                        "meeting": "文字剪貼簿導入",
                        "date": pd.Timestamp.now().strftime("%Y.%m.%d"),
                        "status": "in_progress",
                        "priority": "medium",
                        "content": trimmed,
                        "progress": "自文字段落提取",
                        "is_cross_dept": False
                    })

    if extracted_tasks:
        st.markdown(f"#### 🔎 解析結果：找到 {len(extracted_tasks)} 項可能待辦")
        st.info("請勾選您想加入追蹤的項目，並確認其指派部門與負責人。若發現同名任務，將會自動合併至該任務之歷史會議中。")
        
        import_list = []
        for idx, task in enumerate(extracted_tasks):
            # Run fuzzy match calculation
            best_match, score, reason = find_fuzzy_match(task, st.session_state.tasks)
            
            with st.expander(f"項目 {idx+1}: {task['title']} ({task['dept']})", expanded=True):
                col_sel, col_det = st.columns([1, 9])
                with col_sel:
                    to_import = st.checkbox("匯入此項", value=True, key=f"imp_cb_{idx}")
                with col_det:
                    t_title = st.text_input("任務名稱", value=task["title"], key=f"imp_title_{idx}")
                    col_det_sub1, col_det_sub2 = st.columns(2)
                    with col_det_sub1:
                        t_dept = st.selectbox("主責部門", ["教材研發組", "社工特教組", "學區營運組", "畢業生組", "處長室/行政管理"], index=["教材研發組", "社工特教組", "學區營運組", "畢業生組", "處長室/行政管理"].index(task["dept"]), key=f"imp_dept_{idx}")
                        t_owner = st.text_input("主責負責人", value=task["owner"], key=f"imp_owner_{idx}")
                    with col_det_sub2:
                        t_status = st.selectbox("狀態", ["pending", "in_progress", "completed"], index=["pending", "in_progress", "completed"].index(task["status"]), key=f"imp_status_{idx}")
                        t_priority = st.selectbox("優先度", ["high", "medium", "low"], index=["high", "medium", "low"].index(task["priority"]), key=f"imp_pri_{idx}")
                        t_date = st.text_input("會議日期 (格式: YYYY.MM.DD)", value=task["date"], key=f"imp_date_{idx}")
                    
                    t_cross = st.checkbox("🤝 此任務需要跨部門/跨組別協作配合", value=bool(task.get("is_cross_dept", False)), key=f"imp_cross_{idx}")
                    t_content = st.text_area("決議說明", value=task["content"], key=f"imp_cont_{idx}")
                    t_progress = st.text_input("最新進度", value=task["progress"], key=f"imp_prog_{idx}")
                    
                    # Fuzzy match choices
                    merge_into_id = None
                    if best_match:
                        # Color coding based on similarity score (green for high/exact, orange for medium)
                        color = "#10B981" if score >= 0.85 else "#F59E0B"
                        st.markdown(clean_html(f"""
                        <div style="background: rgba(16,185,129,0.03); border: 1px solid {color}; padding: 10px; border-radius: 6px; font-size: 11.5px; margin-top: 8px; margin-bottom: 8px;">
                            💡 <b>智慧相似度偵測</b>：此項目與現有任務高度相似！({reason})<br>
                            👉 現有任務：<b>{html_escape(best_match['title'])}</b> (主責: {html_escape(best_match['owner'])} | 最近進度: {html_escape(best_match['progress'])})
                        </div>
                        """), unsafe_allow_html=True)
                        
                        # Let user choose whether to merge or keep separate. Default is "directly merge" if high score.
                        default_act_idx = 1 if score >= 0.80 else 0
                        merge_action = st.radio("匯入決策：", ["建立為獨立新任務卡片", f"直接合併併入現有任務：{best_match['title']}"], index=default_act_idx, key=f"imp_merge_act_{idx}")
                        if "直接合併併入現有任務" in merge_action:
                            merge_into_id = best_match['id']
                    
                    if to_import:
                        import_list.append({
                            "id": task["id"],
                            "title": t_title,
                            "dept": t_dept,
                            "owner": t_owner or "待指派",
                            "meeting": task["meeting"],
                            "date": t_date,
                            "status": t_status,
                            "priority": t_priority,
                            "content": t_content,
                            "progress": t_progress,
                            "is_cross_dept": t_cross,
                            "merge_into_id": merge_into_id
                        })
        
        if st.button("📥 匯入已勾選之任務卡片", type="primary"):
            if import_list:
                for item in import_list:
                    existing_task = None
                    
                    # 1. If explicit merge ID is selected
                    if item.get("merge_into_id"):
                        for t in st.session_state.tasks:
                            if t["id"] == item["merge_into_id"]:
                                existing_task = t
                                break
                                
                    # 2. Fallback to exact title match
                    if not existing_task:
                        for t in st.session_state.tasks:
                            if t.get("title", "").strip().lower() == item["title"].strip().lower():
                                existing_task = t
                                break
                    
                    # Track imported tasks
                    imported_task_ref = None
                    if existing_task:
                        if "meeting_history" not in existing_task:
                            existing_task["meeting_history"] = []
                        already_has_meeting = any(mh.get("meeting") == item["meeting"] for mh in existing_task["meeting_history"])
                        if not already_has_meeting:
                            existing_task["meeting_history"].append({
                                "meeting": item["meeting"],
                                "date": item["date"],
                                "content": item["content"]
                            })
                            
                        if "progress_history" not in existing_task:
                            existing_task["progress_history"] = []
                        already_has_progress = any(ph.get("text") == item["progress"] and ph.get("date") == item["date"] for ph in existing_task["progress_history"])
                        if not already_has_progress:
                            existing_task["progress_history"].append({
                                "date": item["date"],
                                "text": item["progress"]
                            })
                            
                        # Chronological descending sort (latest date first)
                        existing_task["meeting_history"].sort(key=lambda x: x.get("date", ""), reverse=True)
                        existing_task["progress_history"].sort(key=lambda x: x.get("date", ""), reverse=True)
                        
                        # Set active fields dynamically from the chronologically latest history entry
                        latest_meeting = existing_task["meeting_history"][0]
                        latest_progress = existing_task["progress_history"][0]
                        
                        existing_task["meeting"] = latest_meeting["meeting"]
                        existing_task["date"] = latest_meeting["date"]
                        existing_task["content"] = latest_meeting["content"]
                        existing_task["progress"] = latest_progress["text"]
                        
                        # Only update status/priority/is_cross if this imported item is the chronologically latest
                        if latest_meeting["date"] == item["date"]:
                            existing_task["status"] = item["status"]
                            existing_task["priority"] = item["priority"]
                            existing_task["is_cross_dept"] = item["is_cross_dept"]
                        imported_task_ref = existing_task
                    else:
                        item_new = {
                            "id": item["id"],
                            "title": item["title"],
                            "dept": item["dept"],
                            "owner": item["owner"],
                            "meeting": item["meeting"],
                            "date": item["date"],
                            "status": item["status"],
                            "priority": item["priority"],
                            "content": item["content"],
                            "progress": item["progress"],
                            "is_cross_dept": item["is_cross_dept"],
                            "created_date": item["date"],
                            "meeting_history": [{
                                "meeting": item["meeting"],
                                "date": item["date"],
                                "content": item["content"]
                            }],
                            "progress_history": [{
                                "date": item["date"],
                                "text": item["progress"]
                            }]
                        }
                        st.session_state.tasks.insert(0, item_new)
                        imported_task_ref = item_new
                    
                    if imported_task_ref and imported_task_ref not in modified_tasks_list:
                        modified_tasks_list.append(imported_task_ref)
                        
                save_tasks(st.session_state.tasks, settings, modified_tasks=modified_tasks_list)
                st.success(f"✅ 成功匯入/更新了 {len(import_list)} 項任務！")
                st.rerun()
            else:
                st.warning("請至少勾選一項進行匯入！")

# ---- TAB 3: 業務分類字典 ----
with tab_dictionary:
    st.markdown("### 📖 各部門業務分類字典")
    st.markdown("本字典定義了教育輔導處各組別的核心業務，系統在匯入新會議紀錄時會自動參照本字典的關鍵字進行分類。")
    dictionary_data = [
        {
            "name": "4.1 處長室 / 跨部門協調",
            "desc": "規劃制定教育輔導相關服務的總體策略，預算與資源分配管理，帶領處內各組達成年度工作目標，督導與跨部門行政事務協調。",
            "items": ["A. 瞭解趨勢，制定服務總體策略", "B. 監督與提供內部績效評估及優化建議", "C. 督導處內組長、教育訓練規劃", "D. 管理預算與資源", "E. 跨部門教育輔導事宜協調與管理"]
        },
        {
            "name": "4.2 教材研發編輯組",
            "desc": "研發與優化博幼教材（英文、數學、閱讀、科普、程式設計），教材版面設計與排版，組織在學生多元學習活動，推廣 PASSION 英文種子教師觀議課機制。",
            "items": ["A. 研發編輯、出版課輔教材與檢測卷", "B. 收集現場教材反饋並改進成效", "C. 教材美編、版面設計與印刷前置", "D. 策劃全機構在學生多元學習活動", "E. 辦理課輔老師英數檢測與培訓認證", "F. 推行 PASSION 英文種子老師觀議課加給"]
        },
        {
            "name": "4.3 社會工作暨特殊教育發展組",
            "desc": "綜理機構社會工作與特殊教育的制度規劃與執行，輔導特殊需求學生成效評估。提供高關懷學生處遇追蹤、國三畢業典禮籌備，管理急難救助金（李醒嘉安心基金）與早療實驗計畫。",
            "items": ["A. 倡議與建置社工督導及增能課程", "B. 提供高關懷學生追蹤處遇與外部資源連結", "C. 特殊需求學生資格評估與輔導成效追蹤", "D. 研發特教生教案設計（生活管理、人際、專注力）", "E. 執行傳善獎早療實驗計畫、早療平台與小手計畫", "F. 統籌國三畢業典禮，審核與發放李醒嘉安心基金"]
        },
        {
            "name": "4.4 學區營運發展組",
            "desc": "綜理課輔中心學區營運、課輔老師招募及考評調薪。追蹤與分析學區教學品質、學生會考及檢測成績，系統開發（學生與成績系統），以及建構數位學習平台（均一）任務學習模式與手冊大綱。",
            "items": ["A. 制定課輔中心學區運作制度與課師培訓考核", "B. 追蹤分析教學品質與學生檢測通過率、會考表現", "C. 教育輔導處資訊系統（成績/學生資料）評估維護", "D. 優化線上及紙本週誌填寫，降低課師行政壓力", "E. 執行尖前中心數位平台（均一）任務學習實驗計畫與手冊大綱"]
        },
        {
            "name": "4.5 畢業生追蹤服務規劃組",
            "desc": "綜理博幼大專與高中職畢業生輔導、升學與就業規劃，申辦與發放各大企業與內部生活助學金，定期追蹤在校成績，並組織實施寒暑假返鄉志工服務機制。",
            "items": ["A. 規劃與管理畢業生高關懷制度，協助評估與輔導", "B. 經營畢業生組織（學生會）運作、辦理幹訓與活動", "C. 追蹤畢業生高中職、大專端就學成績與適應情況", "D. 統籌各大專與高中職獎助學金申請、審核與分攤核銷", "E. 審查補習/家教申請，媒合企業職缺與友善企業合作", "F. 畢業生翻譯培力、寒暑假回娘家活動統籌"]
        }
    ]

    for dept in dictionary_data:
        st.markdown(f"#### 👤 {dept['name']}")
        st.markdown(clean_html(f"<div style='font-size: 12px; color: var(--text-muted); background: var(--card-bg); padding: 10px; border-radius: 6px; border: 1px solid var(--border-color); margin-bottom: 8px;'>{dept['desc']}</div>"), unsafe_allow_html=True)
        cols = st.columns(3)
        for i, item in enumerate(dept['items']):
            with cols[i % 3]:
                st.markdown(clean_html(f"<div style='font-size: 11px; padding: 4px 8px; margin-bottom: 4px; border-left: 2px solid #10B981; color: var(--text);'>{item}</div>"), unsafe_allow_html=True)
        st.markdown("---")

# ---- TAB 4: 歷史歸檔檔案庫 (Scheme A implementation) ----
with tab_archive:
    st.markdown("### 📦 歷史歸檔檔案庫")
    st.markdown("此處列出所有已完成並手動歸檔的歷史任務項目。您可以在此查閱決議與歷史進度，或將其重新喚醒至主看板。")
    
    arc_search = st.text_input("搜尋已歸檔任務名稱、負責人或決議說明...", key="arc_search_input")
    arc_dept = st.selectbox("篩選主責部門：", ["全部", "教材研發組", "社工特教組", "學區營運組", "畢業生組", "處長室/行政管理"], key="arc_dept_select")
    
    # Query archived tasks (includes archived tasks, plus completed tasks older than 3 months)
    import pandas as pd
    archived_tasks = []
    for t in tasks:
        if t.get("status") == "archived":
            archived_tasks.append(t)
        elif t.get("status") == "completed":
            try:
                task_date = pd.to_datetime(t.get("date", ""), format="%Y.%m.%d")
                three_months_ago = pd.Timestamp.now() - pd.DateOffset(months=3)
                if task_date < three_months_ago:
                    archived_tasks.append(t)
            except:
                pass
    
    # Filter archived tasks
    if arc_search:
        archived_tasks = [t for t in archived_tasks if arc_search.lower() in t.get("title", "").lower() or arc_search.lower() in t.get("owner", "").lower() or arc_search.lower() in t.get("content", "").lower()]
    if arc_dept != "全部":
        archived_tasks = [t for t in archived_tasks if t.get("dept") == arc_dept]
        
    if not archived_tasks:
        st.info("目前沒有符合篩選條件的歸檔歷史任務。")
    else:
        # Table Grid Headers
        a_cols = st.columns([1.2, 1.5, 4.0, 1.2, 3.6, 0.5])
        a_cols[0].markdown("**狀態**")
        a_cols[1].markdown("**主責部門**")
        a_cols[2].markdown("**任務項目 / 決議說明**")
        a_cols[3].markdown("**負責人**")
        a_cols[4].markdown("**最新執行進度**")
        a_cols[5].markdown("**喚醒**")
        st.markdown("<hr style='margin: 0.5rem 0 1rem 0; border-color: var(--border-color);'>", unsafe_allow_html=True)
        
        # Grid Data Rows for Archived Tasks
        for idx, t in enumerate(archived_tasks):
            esc_title = html_escape(t.get("title", ""))
            esc_content = html_escape(t.get("content", ""))
            esc_progress = html_escape(t.get("progress", ""))
            esc_owner = html_escape(t.get("owner", ""))
            esc_dept = html_escape(t.get("dept", ""))
            esc_date = html_escape(t.get("date", ""))

            # Calculate project duration
            start_date = get_task_start_date(t)
            period_str = f"{start_date} ~ {esc_date} (已完結)"

            r_cols = st.columns([1.2, 1.5, 4.0, 1.2, 3.6, 0.5])
            
            if t.get("status") == "completed":
                r_cols[0].markdown('<span class="badge" style="background: rgba(16,185,129,0.15); color: #10B981; border: 1px solid rgba(16,185,129,0.15);">&#9989; &#27511;&#21490;&#23436;&#25104;</span>', unsafe_allow_html=True)
            else:
                r_cols[0].markdown('<span class="badge" style="background: rgba(148,163,184,0.15); color: #475569; border: 1px solid rgba(148,163,184,0.25);">&#128230; &#24050;&#27512;&#27284;</span>', unsafe_allow_html=True)
            
            # Column 2: Department Badge
            r_cols[1].markdown(f'<span class="badge" style="background: rgba(99,102,241,0.08); color: #6366F1; border: 1px solid rgba(99,102,241,0.15);">{esc_dept}</span>', unsafe_allow_html=True)
            
            # Column 3: Task & description & meeting history dropdown
            mh_list = t.get("meeting_history", [])
            meeting_history_html = "".join([f'<li style="margin-bottom:3px;"><b>{html_escape(mh.get("date"))}</b>: {html_escape(mh.get("meeting"))} - <i>{html_escape(mh.get("content"))}</i></li>' for mh in mh_list])
            
            r_cols[2].markdown(clean_html(f"""
            <div style="font-weight: 700; font-size: 13px; text-decoration: line-through; color: var(--text-muted);">{esc_title}</div>
            <div style="font-size: 11px; color: var(--text-muted); margin-top: 2px; line-height:1.4;">{esc_content}</div>
            <div style="font-size: 10.5px; color: var(--text-muted); margin-top: 3px; font-weight: 500;">📅 期程: {period_str}</div>
            <details style="cursor: pointer; margin-top: 4px;">
                <summary style="font-size: 10px; color: #06B6D4; font-weight:600;">📂 提及歷史 ({len(mh_list)}次)</summary>
                <ul style="margin: 4px 0 0 10px; padding-left: 10px; font-size: 10.5px; color: var(--text-muted); line-height:1.45;">
                    {meeting_history_html}
                </ul>
            </details>
            """), unsafe_allow_html=True)
            
            # Column 4: Owner
            r_cols[3].markdown(f"<div style='font-size: 12px; color: var(--text-muted); padding-top:4px;'>{esc_owner}</div>", unsafe_allow_html=True)
            
            # Column 5: Progress & progress history dropdown
            ph_list = t.get("progress_history", [])
            progress_history_html = "".join([f'<div style="margin-bottom:4px; padding-bottom:4px; border-bottom:1px dashed var(--border-color);"><b>{html_escape(ph.get("date"))}</b>: {html_escape(ph.get("text"))}</div>' for ph in ph_list])
            
            r_cols[4].markdown(clean_html(f"""
            <div style="font-size: 12px; color: var(--text-muted); padding-top:4px;">{esc_progress}</div>
            <details style="cursor: pointer; margin-top: 4px;">
                <summary style="font-size: 10px; color: #10B981; font-weight:600;">📈 歷史進度 ({len(ph_list)}條)</summary>
                <div style="margin: 4px 0 0 0; padding-left: 6px; border-left: 2px solid #10B981; font-size: 10.5px; color: var(--text-muted); line-height:1.45;">
                    {progress_history_html}
                </div>
            </details>
            """), unsafe_allow_html=True)
            
            # Column 6: Action (Wake up / Unarchive)
            with r_cols[5]:
                if st.button("↩️", key=f"unarc_{t['id']}_{idx}", help="喚醒任務並移回進行中", use_container_width=True):
                    t["status"] = "in_progress"
                    t["progress"] = "狀態更新：自歷史檔案庫重啟任務"
                    t["date"] = pd.Timestamp.now().strftime("%Y.%m.%d")
                    t["progress_history"].insert(0, {
                        "date": t["date"],
                        "text": "狀態更新：自歷史檔案庫重啟任務"
                    })
                    save_tasks(st.session_state.tasks, settings, modified_tasks=[t])
                    st.success("✅ 任務已成功移回「進行中」看板！")
                    st.rerun()
            
            # Divider between rows
            st.markdown("<hr style='margin: 0.5rem 0; border-color: var(--border-color); opacity: 0.5;'>", unsafe_allow_html=True)

# ---- TAB 5: 雲端同步設定 ----
with tab_sync:
    st.markdown("### ⚙️ 雲端同步與 Google Sheets 設定")
    st.markdown("本系統支援將所有待辦任務與進度即時備份同步至您的 Google 試算表 (Google Sheets)。當多個主管在雲端編輯時，將使用此通道實現即時同步。")
    
    st.info("💡 **本地備份提示**：即使未設定 Google Sheets，系統也會將資料自動備份在本地的 `tasks.json`。因為此檔案在您的 Google 雲端硬碟同步夾中，它依然會被自動同步。")

    st.markdown("#### 1. 設定 Google 試算表共用連結 (讀取)")
    sheet_url_input = st.text_input(
        "請貼上您的 Google 試算表共用網址（權限請設為：知道連結的任何人均可檢視/編輯）：",
        value=settings.get("sheet_url", ""),
        placeholder="https://docs.google.com/spreadsheets/d/SpreadsheetId/edit?usp=sharing"
    )
    
    st.markdown("#### 2. 設定 Google Apps Script Web App 網址 (寫入)")
    gas_url_input = st.text_input(
        "請貼上您的 Google Apps Script 部署之 Web App 網址（用於將網頁更新回傳寫入試算表）：",
        value=settings.get("gas_url", ""),
        placeholder="https://script.google.com/macros/s/GAS_Deployment_ID/exec"
    )

    if st.button("💾 儲存並套用設定", type="primary"):
        settings["sheet_url"] = sheet_url_input.strip()
        settings["gas_url"] = gas_url_input.strip()
        save_settings(settings)
        st.success("✅ 設定已更新！系統將在下次同步時套用新網址。")
        refresh_data()

    st.markdown("---")
    st.markdown("#### 📤 初始化雲端資料庫")
    st.markdown("如果您的 Google 試算表目前是全新的空白表格，您可以點擊下方按鈕，將目前本機已有的所有任務與進度一次性推送上傳到 Google 試算表中，完成資料庫的初始化。")
    if st.button("📤 強制推送本地資料庫至雲端試算表", use_container_width=True):
        if settings.get("gas_url"):
            with st.spinner("正在上傳並同步本地資料至雲端試算表..."):
                save_tasks(st.session_state.tasks, settings)
                st.success("✅ 成功將本地資料推送至雲端試算表！請重新整理您的 Google 試算表網頁查看資料。")
        else:
            st.error("❌ 請先設定並儲存 Google Apps Script Web App 網址！")

    with st.expander("🛠️ 如何設定 Google Apps Script 以支援協同寫入？ (2分鐘教學)", expanded=False):
        st.markdown("""
        **步驟一：**
        打開您的 Google 試算表，點選上方選單的 **「擴充功能」 ➔ 「Apps Script」**。
        
        **步驟二：**
        清空編輯器，並複製貼上以下代碼：
        ```javascript
        function doPost(e) {
          var sheet = SpreadsheetApp.getActiveSpreadsheet().getActiveSheet();
          sheet.clearContents(); // 清空舊內容
          
          var tasks = JSON.parse(e.postData.contents);
          if (tasks.length > 0) {
            // 寫入標題列
            var headers = Object.keys(tasks[0]);
            sheet.appendRow(headers);
            
            // 寫入資料列
            for (var i = 0; i < tasks.length; i++) {
              var row = [];
              for (var j = 0; j < headers.length; j++) {
                row.push(tasks[i][headers[j]] || "");
              }
              sheet.appendRow(row);
            }
          }
          return ContentService.createTextOutput("SUCCESS").setMimeType(ContentService.MimeType.TEXT);
        }
        ```
        
        **步驟三：**
        點選右上角 **「部署」 ➔ 「新增部署」**。
        - 類型選擇 **「網頁應用程式」**。
        - 專案負責人：選您的帳號。
        - 誰有權限存取：選 **「任何人」 (Anyone)** ➔ 這一步非常重要！
        - 點選「部署」並授權，複製產生的 **「網頁應用程式網址」** (即為上面的第2項網址)。
        """)
